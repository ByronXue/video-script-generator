
import streamlit as st
import base64
import io
from docx import Document
from reportlab.pdfgen import canvas
from utils import generate_script

# 标题
st.title("🎬️视频脚本生成器")

# 侧边栏
with st.sidebar:
    # 模型选择下拉框
    model_choice = st.selectbox(
        "请选择API模型",
        options=[
            "deepseek-chat",
            "gpt-3.5-turbo",
            "gpt-4",
        ],
        index=0
    )

    # 根据选择的模型显示相应的获取链接
    if model_choice == "deepseek-chat":
        st.markdown("[获取DeepSeek API秘钥](https://platform.deepseek.com/)")
    elif model_choice in ["gpt-3.5-turbo", "gpt-4"]:
        st.markdown("[获取OpenAI API秘钥](https://platform.openai.com/api-keys)")

    # API密钥输入
    api_key = st.text_input("请输入API秘钥", type="password")

# 主体
subject = st.text_input("💡请输入视频主题")
video_length = st.number_input("⏱️请输入视频的大致时长(单位:分钟)", min_value=0.1, step=0.1)
creativity = st.slider("🤔请选择视频的创造力(数字越小越严谨,反之更多样)", min_value=0.1, max_value=1.0, value=.5, step=0.1)
submit = st.button("📽️生成视频脚本")

# 初始化 session_state 使用 st.session_state 持久化生成结果
if "script_data" not in st.session_state:
    st.session_state.script_data = None

# 生成逻辑
if submit and api_key and subject:
    try:
        with st.spinner("视频脚本生成中..."):
            search_result, title, script = generate_script(subject, video_length, creativity, api_key, model_choice)
        st.session_state.script_data = {
            "title": title,
            "script": script,
            "subject": subject,
            "search_result": search_result
        }
    except Exception as e:
        st.error(f"❌ 生成失败：{type(e).__name__}: {str(e)}")
        st.code(str(e), language="text")
        st.session_state.script_data = None

# 渲染结果（只要 script_data 存在就显示）
if st.session_state.script_data:
    data = st.session_state.script_data
    st.success("✅ 视频脚本已生成！")
    st.subheader("🔥标题:")
    st.write(data["title"])
    st.subheader("📝视频脚本:")
    st.write(data["script"])

    # 下载区域
    st.subheader("📤 脚本下载")

    # TXT
    txt_content = f"标题：{data['title']}\n\n脚本：\n{data['script']}".encode()
    st.download_button(
        label="📄 下载 TXT (.txt)",
        data=txt_content,
        file_name=f"script_{data['subject'].replace(' ', '_')[:20]}.txt",
        mime="text/plain",
        key="btn_txt"
    )

    # Markdown
    md_content = f"# {data['title']}\n\n{data['script']}".encode()
    st.download_button(
        label="📝 下载 Markdown (.md)",
        data=md_content,
        file_name=f"script_{data['subject'].replace(' ', '_')[:20]}.md",
        mime="text/markdown",
        key="btn_md"
    )

    # DOCX
    doc = Document()
    doc.add_heading(data["title"], level=1)
    doc.add_paragraph(data["script"])
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    st.download_button(
        label="📄 下载 DOCX (.docx)",
        data=buffer.getvalue(),
        file_name=f"script_{data['subject'].replace(' ', '_')[:20]}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="btn_docx"
    )

    # 维基百科结果
    with st.expander("维基百科搜索结果 👀"):
        st.write(data["search_result"])